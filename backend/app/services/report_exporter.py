"""
Report Exporter — converts markdown reports to PDF and DOCX formats.
"""

import re
import tempfile
from typing import List, Tuple

import fitz  # PyMuPDF
# python-docx is imported lazily inside export_to_docx() to avoid blocking PDF export
# if the package is not installed.

from ..utils.logger import get_logger

logger = get_logger('mirofish.report_exporter')


# ── Markdown Parser ──

def _parse_markdown_blocks(md: str) -> List[Tuple[str, str]]:
    """
    Parse markdown into a list of (block_type, content) tuples.

    Block types: h1, h2, h3, quote, ul, ol, hr, paragraph
    """
    blocks = []
    lines = md.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if re.match(r'^-{3,}$|^\*{3,}$|^_{3,}$', stripped):
            blocks.append(('hr', ''))
            i += 1
            continue

        # Headings
        if stripped.startswith('#### '):
            blocks.append(('h4', stripped[5:].strip()))
            i += 1
            continue
        if stripped.startswith('### '):
            blocks.append(('h3', stripped[4:].strip()))
            i += 1
            continue
        if stripped.startswith('## '):
            blocks.append(('h2', stripped[3:].strip()))
            i += 1
            continue
        if stripped.startswith('# '):
            blocks.append(('h1', stripped[2:].strip()))
            i += 1
            continue

        # Blockquote (may span multiple lines)
        if stripped.startswith('>'):
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith('>'):
                quote_lines.append(lines[i].strip().lstrip('>').strip())
                i += 1
            blocks.append(('quote', ' '.join(quote_lines)))
            continue

        # Unordered list
        if re.match(r'^[-*+]\s', stripped):
            list_items = []
            while i < len(lines) and re.match(r'^\s*[-*+]\s', lines[i]):
                list_items.append(re.sub(r'^\s*[-*+]\s+', '', lines[i]).strip())
                i += 1
            blocks.append(('ul', list_items))
            continue

        # Ordered list
        if re.match(r'^\d+\.\s', stripped):
            list_items = []
            while i < len(lines) and re.match(r'^\s*\d+\.\s', lines[i]):
                list_items.append(re.sub(r'^\s*\d+\.\s+', '', lines[i]).strip())
                i += 1
            blocks.append(('ol', list_items))
            continue

        # Regular paragraph (collect consecutive non-special lines)
        para_lines = []
        while i < len(lines):
            l = lines[i].strip()
            if not l or l.startswith('#') or l.startswith('>') or re.match(r'^[-*+]\s', l) or re.match(r'^\d+\.\s', l) or re.match(r'^-{3,}$|^\*{3,}$|^_{3,}$', l):
                break
            para_lines.append(l)
            i += 1
        if para_lines:
            blocks.append(('paragraph', ' '.join(para_lines)))

    return blocks


def _split_inline(text: str) -> List[Tuple[str, str]]:
    """
    Split text into inline segments: ('bold', text), ('italic', text), ('normal', text).
    Handles **bold** and *italic* / _italic_ markers.
    """
    segments = []
    pattern = re.compile(r'\*\*(.+?)\*\*|\*(.+?)\*|_(.+?)_')
    last = 0

    for m in pattern.finditer(text):
        if m.start() > last:
            segments.append(('normal', text[last:m.start()]))
        if m.group(1):
            segments.append(('bold', m.group(1)))
        elif m.group(2):
            segments.append(('italic', m.group(2)))
        elif m.group(3):
            segments.append(('italic', m.group(3)))
        last = m.end()

    if last < len(text):
        segments.append(('normal', text[last:]))

    return segments


# ── PDF Export (PyMuPDF) ──

def export_to_pdf(markdown_content: str, title: str = "Report") -> str:
    """
    Convert markdown to a formatted PDF.

    Args:
        markdown_content: Full report markdown
        title: Report title (used in header)

    Returns:
        Path to generated PDF temp file
    """
    blocks = _parse_markdown_blocks(markdown_content)

    doc = fitz.open()
    page_width = 595  # A4
    page_height = 842
    margin_left = 60
    margin_right = 60
    margin_top = 60
    margin_bottom = 60
    usable_width = page_width - margin_left - margin_right

    # Fonts
    font_regular = "helv"
    font_bold = "hebo"
    font_italic = "heit"

    page = doc.new_page(width=page_width, height=page_height)
    y = margin_top

    def _new_page():
        nonlocal page, y
        # Page number on current page
        _draw_page_number(page, doc.page_count)
        page = doc.new_page(width=page_width, height=page_height)
        y = margin_top

    def _draw_page_number(pg, num):
        footer_text = f"— {num} —"
        tw = fitz.get_text_length(footer_text, fontname=font_regular, fontsize=8)
        pg.insert_text(
            (page_width / 2 - tw / 2, page_height - 30),
            footer_text, fontname=font_regular, fontsize=8,
            color=(0.5, 0.5, 0.5)
        )

    def _check_space(needed):
        nonlocal y
        if y + needed > page_height - margin_bottom:
            _new_page()

    def _write_text(text, fontname, fontsize, color=(0, 0, 0), indent=0, spacing=4):
        nonlocal y
        max_width = usable_width - indent
        # Wrap text manually
        words = text.split(' ')
        lines_out = []
        current_line = ""
        for word in words:
            test = (current_line + " " + word).strip()
            tw = fitz.get_text_length(test, fontname=fontname, fontsize=fontsize)
            if tw > max_width and current_line:
                lines_out.append(current_line)
                current_line = word
            else:
                current_line = test
        if current_line:
            lines_out.append(current_line)

        line_height = fontsize * 1.4
        for line_text in lines_out:
            _check_space(line_height)
            page.insert_text(
                (margin_left + indent, y + fontsize),
                line_text, fontname=fontname, fontsize=fontsize,
                color=color
            )
            y += line_height
        y += spacing

    def _write_inline_text(text, base_fontsize, indent=0, spacing=4):
        """Write text with inline bold/italic formatting."""
        nonlocal y
        segments = _split_inline(text)
        # For simplicity, render each segment sequentially
        # Build full plain text for wrapping, then render with formatting
        plain = text.replace('**', '').replace('*', '').replace('_', '')
        max_width = usable_width - indent

        # Simple approach: render with primary font, bold segments get bold font
        words = []
        for style, seg_text in segments:
            for w in seg_text.split(' '):
                if w:
                    words.append((style, w))

        lines_out = []
        current_line_words = []
        current_width = 0

        for style, word in words:
            fn = font_bold if style == 'bold' else (font_italic if style == 'italic' else font_regular)
            ww = fitz.get_text_length(word + ' ', fontname=fn, fontsize=base_fontsize)
            if current_width + ww > max_width and current_line_words:
                lines_out.append(list(current_line_words))
                current_line_words = [(style, word)]
                current_width = ww
            else:
                current_line_words.append((style, word))
                current_width += ww

        if current_line_words:
            lines_out.append(current_line_words)

        line_height = base_fontsize * 1.4
        for line_words in lines_out:
            _check_space(line_height)
            x = margin_left + indent
            for style, word in line_words:
                fn = font_bold if style == 'bold' else (font_italic if style == 'italic' else font_regular)
                page.insert_text(
                    (x, y + base_fontsize),
                    word + ' ', fontname=fn, fontsize=base_fontsize,
                    color=(0, 0, 0)
                )
                x += fitz.get_text_length(word + ' ', fontname=fn, fontsize=base_fontsize)
            y += line_height
        y += spacing

    # Render blocks
    for block_type, content in blocks:
        if block_type == 'h1':
            _check_space(40)
            y += 10
            _write_text(content, font_bold, 20, color=(0.1, 0.1, 0.1), spacing=12)
            # Draw underline
            page.draw_line(
                fitz.Point(margin_left, y - 6),
                fitz.Point(page_width - margin_right, y - 6),
                color=(0.7, 0.7, 0.7), width=0.5
            )
            y += 4

        elif block_type == 'h2':
            _check_space(30)
            y += 8
            _write_text(content, font_bold, 15, color=(0.15, 0.15, 0.15), spacing=8)

        elif block_type in ('h3', 'h4'):
            _check_space(24)
            y += 4
            _write_text(content, font_bold, 12, color=(0.2, 0.2, 0.2), spacing=6)

        elif block_type == 'quote':
            _check_space(20)
            # Draw left border
            quote_start_y = y
            _write_inline_text(content, 10, indent=20, spacing=2)
            # Italic color for quotes
            page.draw_line(
                fitz.Point(margin_left + 8, quote_start_y),
                fitz.Point(margin_left + 8, y - 2),
                color=(0.3, 0.5, 0.8), width=2
            )
            y += 6

        elif block_type == 'ul':
            for item in content:
                _check_space(16)
                _write_inline_text(f"  \u2022  {item}", 11, indent=10, spacing=2)
            y += 4

        elif block_type == 'ol':
            for idx, item in enumerate(content, 1):
                _check_space(16)
                _write_inline_text(f"  {idx}.  {item}", 11, indent=10, spacing=2)
            y += 4

        elif block_type == 'hr':
            _check_space(20)
            y += 8
            page.draw_line(
                fitz.Point(margin_left, y),
                fitz.Point(page_width - margin_right, y),
                color=(0.8, 0.8, 0.8), width=0.5
            )
            y += 12

        elif block_type == 'paragraph':
            _check_space(16)
            _write_inline_text(content, 11, spacing=6)

    # Final page number
    _draw_page_number(page, doc.page_count)

    # Save — close temp file first so PyMuPDF can write to it on Windows
    tmp = tempfile.NamedTemporaryFile(suffix='.pdf', delete=False)
    tmp_path = tmp.name
    tmp.close()
    page_count = doc.page_count
    doc.save(tmp_path)
    doc.close()
    logger.info(f"PDF exported: {tmp_path} ({page_count} pages)")
    return tmp_path


# ── DOCX Export (python-docx) ──

def export_to_docx(markdown_content: str, title: str = "Report") -> str:
    """
    Convert markdown to a formatted DOCX document.

    Args:
        markdown_content: Full report markdown
        title: Report title

    Returns:
        Path to generated DOCX temp file
    """
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor

    blocks = _parse_markdown_blocks(markdown_content)

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Configure heading styles
    for level, size in [(1, 20), (2, 16), (3, 13)]:
        heading_style = doc.styles[f'Heading {level}']
        heading_style.font.size = Pt(size)
        heading_style.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)

    def _add_inline_runs(paragraph, text):
        """Add text with inline bold/italic formatting to a paragraph."""
        segments = _split_inline(text)
        for seg_style, seg_text in segments:
            run = paragraph.add_run(seg_text)
            if seg_style == 'bold':
                run.bold = True
            elif seg_style == 'italic':
                run.italic = True

    for block_type, content in blocks:
        if block_type == 'h1':
            doc.add_heading(content, level=1)

        elif block_type == 'h2':
            doc.add_heading(content, level=2)

        elif block_type in ('h3', 'h4'):
            doc.add_heading(content, level=3)

        elif block_type == 'quote':
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(content)
            run.italic = True
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

        elif block_type == 'ul':
            for item in content:
                p = doc.add_paragraph(style='List Bullet')
                _add_inline_runs(p, item)

        elif block_type == 'ol':
            for item in content:
                p = doc.add_paragraph(style='List Number')
                _add_inline_runs(p, item)

        elif block_type == 'hr':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            # Thin line via border
            run = p.add_run('─' * 60)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            run.font.size = Pt(8)

        elif block_type == 'paragraph':
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            _add_inline_runs(p, content)

    # Save — close temp file first so python-docx can write to it on Windows
    tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    tmp_path = tmp.name
    tmp.close()
    doc.save(tmp_path)
    logger.info(f"DOCX exported: {tmp_path}")
    return tmp_path
